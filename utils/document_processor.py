"""Document processor for handling different file types."""
from io import BytesIO
import tempfile

from pypdf import PdfReader
from docx import Document
from PIL import Image


class DocumentProcessor:
    """Process different types of documents."""

    def __init__(self):
        self.supported_formats = {
            'pdf': self.process_pdf,
            'doc': self.process_doc,
            'docx': self.process_doc,
            'png': self.process_image,
            'jpg': self.process_image,
            'jpeg': self.process_image,
        }

    def get_file_type(self, filename: str) -> str:
        """Extract file extension."""
        return filename.split('.')[-1].lower()

    def is_supported(self, filename: str) -> bool:
        """Check if file type is supported."""
        file_type = self.get_file_type(filename)
        return file_type in self.supported_formats

    def process_file(self, uploaded_file) -> dict:
        """Process uploaded file based on its type."""
        file_type = self.get_file_type(uploaded_file.name)

        if not self.is_supported(uploaded_file.name):
            raise ValueError(f"Unsupported file type: {file_type}")

        processor = self.supported_formats[file_type]
        return processor(uploaded_file)

    def process_pdf(self, uploaded_file) -> dict:
        """Extract text from PDF."""
        pdf_reader = PdfReader(BytesIO(uploaded_file.read()))

        text_content = []
        for page_num, page in enumerate(pdf_reader.pages):
            text = page.extract_text()
            if text.strip():
                text_content.append({
                    'page': page_num + 1,
                    'content': text
                })

        return {
            'type': 'pdf',
            'filename': uploaded_file.name,
            'pages': text_content,
            'total_pages': len(pdf_reader.pages)
        }

    def process_doc(self, uploaded_file) -> dict:
        """Extract text from DOCX."""
        doc = Document(BytesIO(uploaded_file.read()))

        text_content = []
        for para_num, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                text_content.append({
                    'paragraph': para_num + 1,
                    'content': paragraph.text
                })

        return {
            'type': 'docx',
            'filename': uploaded_file.name,
            'paragraphs': text_content,
            'total_paragraphs': len(doc.paragraphs)
        }

    def process_image(self, uploaded_file) -> dict:
        """Process image file."""
        # Save image temporarily for processing
        image = Image.open(BytesIO(uploaded_file.read()))

        # Create a temporary file to store the image
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{self.get_file_type(uploaded_file.name)}") as tmp_file:
            image.save(tmp_file.name)
            temp_path = tmp_file.name

        return {
            'type': 'image',
            'filename': uploaded_file.name,
            'temp_path': temp_path,
            'size': image.size,
            'format': image.format
        }
